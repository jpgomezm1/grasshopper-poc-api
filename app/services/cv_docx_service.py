"""La hoja de vida en Word · segundo renderizador desde el mismo `CVData`.

## Por qué existe

El PDF es un documento cerrado. A un estudiante que se está postulando le sirve
poder **abrir el archivo y ajustar una frase** antes de mandarlo, y muchas
convocatorias piden explícitamente `.doc`/`.docx`.

Además resuelve un problema que el PDF tiene y que no es del estudiante: la
generación del PDF **depende de GTK y no corre en Windows**, así que en local no
se puede verificar. Esto es Python puro — se genera igual en todas partes.

## La regla que comparte con el PDF

`CVData` sigue siendo el único modelo de contenido y `cv_variants` sigue siendo
la única política: **el estándar decide qué sale y en qué orden**, aquí también.
Si `us` omite la foto en el PDF y no la omitiera aquí, el estudiante descargaría
dos documentos distintos creyendo que son el mismo — que es justo el tipo de
incoherencia que este repo ya pagó una vez (P0-8: dos sitios decidiendo lo mismo).
Por eso la decisión se importa, no se reimplementa.
"""
from __future__ import annotations

import base64
import io
import logging
from typing import Any, List, Optional

from app.services.cv_pdf_service import CVData
from app.services.cv_variants import (
    ESTANDAR_POR_DEFECTO,
    ESTILO_POR_DEFECTO,
    debe_incluir_foto,
    idioma_va_en_seccion,
    obtener_estandar,
    obtener_estilo,
)

logger = logging.getLogger(__name__)

# Los mismos de la marca que usa el CSS del PDF.
_NAVY = (0x16, 0x41, 0x94)
_LIMA_OSCURO = (0x5B, 0x6B, 0x00)
_GRIS = (0x5B, 0x64, 0x70)


def _rgb(color):
    from docx.shared import RGBColor

    return RGBColor(*color)


def _foto_bytes(data_uri: Optional[str]) -> Optional[bytes]:
    """Saca los bytes del data URI · devuelve None si viene mal formado.

    Una foto corrupta no puede tumbar la descarga entera: la hoja de vida sin
    foto sigue sirviendo, una excepción no.
    """
    if not data_uri or "," not in data_uri:
        return None
    try:
        return base64.b64decode(data_uri.split(",", 1)[1])
    except Exception:  # noqa: BLE001
        logger.warning("Foto en data URI ilegible · el CV sale sin foto")
        return None


def _titulo(doc, texto: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(texto.upper())
    run.bold = True
    run.font.color.rgb = _rgb(_NAVY)
    run.font.size = _pt(12)
    p.paragraph_format.space_before = _pt(10)
    p.paragraph_format.space_after = _pt(3)


def _pt(n):
    from docx.shared import Pt

    return Pt(n)


def _linea_lista(doc, etiqueta: str, valores: List[str], color) -> None:
    if not valores:
        return
    p = doc.add_paragraph()
    run = p.add_run(f"{etiqueta}: ")
    run.bold = True
    run.font.color.rgb = _rgb(color)
    run.font.size = _pt(10)
    resto = p.add_run(" · ".join(str(v) for v in valores))
    resto.font.size = _pt(10)


def _seccion_perfil(doc, cv: CVData) -> None:
    if not (cv.summary or cv.strengths or cv.interests or cv.values):
        return
    _titulo(doc, "Perfil")
    if cv.summary:
        p = doc.add_paragraph(cv.summary)
        p.runs[0].font.size = _pt(10)
    _linea_lista(doc, "Fortalezas", cv.strengths, _LIMA_OSCURO)
    _linea_lista(doc, "Áreas de interés", cv.interests, _NAVY)
    _linea_lista(doc, "Valores", cv.values, _LIMA_OSCURO)


def _seccion_idiomas(doc, cv: CVData) -> None:
    """Espejo de `cv_pdf_service._html_idiomas` · si divergen, el estudiante se
    baja dos documentos distintos creyendo que son el mismo (P0-8)."""
    if not cv.english_level:
        return
    _titulo(doc, "Idiomas")
    p = doc.add_paragraph()
    r = p.add_run(f"Inglés · nivel {cv.english_level}")
    r.font.size = _pt(10)
    marco = p.add_run(" (Marco Común Europeo de Referencia · MCER/CEFR)")
    marco.font.size = _pt(9)
    marco.font.color.rgb = _rgb(_GRIS)


def _seccion_tests(doc, cv: CVData) -> None:
    if not cv.test_highlights:
        return
    _titulo(doc, "Resultados de tests")
    # Se desempaqueta por posición · un CVData guardado antes de A3 puede traer
    # tuplas de 3 en vez de 4 (misma razón que en `cv_pdf_service`).
    for fila in cv.test_highlights:
        p = doc.add_paragraph(style="List Bullet")
        etiqueta = p.add_run(f"{fila[0]}: ")
        etiqueta.bold = True
        etiqueta.font.size = _pt(10)
        valor = p.add_run(str(fila[1]))
        valor.bold = True
        valor.font.color.rgb = _rgb(_LIMA_OSCURO)
        valor.font.size = _pt(10)
        if len(fila) > 2 and fila[2]:
            desc = p.add_run(f" — {fila[2]}")
            desc.font.size = _pt(9)
            desc.font.color.rgb = _rgb(_GRIS)


def _seccion_actividades(doc, cv: CVData) -> None:
    _titulo(doc, "Actividades extracurriculares")
    if not cv.activities:
        p = doc.add_paragraph(
            "Aún no hay actividades registradas. Agrégalas en \"Mis actividades\" "
            "para enriquecer tu CV."
        )
        p.runs[0].italic = True
        p.runs[0].font.color.rgb = _rgb(_GRIS)
        p.runs[0].font.size = _pt(9.5)
        return

    for a in cv.activities:
        cat = doc.add_paragraph()
        run_cat = cat.add_run(a.category_label.upper())
        run_cat.bold = True
        run_cat.font.size = _pt(8.5)
        run_cat.font.color.rgb = _rgb(_NAVY)
        cat.paragraph_format.space_before = _pt(6)
        cat.paragraph_format.space_after = _pt(0)

        cabecera = doc.add_paragraph()
        nombre = cabecera.add_run(a.name)
        nombre.bold = True
        nombre.font.size = _pt(10.5)
        meta = [x for x in (a.period, f"{a.hours_per_week} h/sem" if a.hours_per_week else None) if x]
        if meta:
            run_meta = cabecera.add_run(f"   {' · '.join(meta)}")
            run_meta.font.size = _pt(9)
            run_meta.font.color.rgb = _rgb(_GRIS)
        cabecera.paragraph_format.space_after = _pt(0)

        for texto, tam, color in (
            (a.role, 9.5, _GRIS),
            (a.description, 9.5, None),
        ):
            if texto:
                p = doc.add_paragraph(texto)
                p.runs[0].font.size = _pt(tam)
                if color:
                    p.runs[0].font.color.rgb = _rgb(color)
                p.paragraph_format.space_after = _pt(0)

        for logro in a.achievements or []:
            if logro:
                p = doc.add_paragraph(str(logro), style="List Bullet")
                p.runs[0].font.size = _pt(9.5)


_SECCIONES = {
    "perfil": _seccion_perfil,
    "idiomas": _seccion_idiomas,
    "tests": _seccion_tests,
    "actividades": _seccion_actividades,
}


def _encabezado(doc, cv: CVData, *, con_foto: bool, omitir_idioma: bool = False) -> None:
    from docx.shared import Mm

    foto = _foto_bytes(cv.photo_data_uri) if con_foto else None

    # Con foto, el encabezado va en una tabla invisible de 2 columnas: es la
    # única forma fiable de poner texto e imagen lado a lado en Word.
    destino = doc
    primer_parrafo = None
    if foto:
        tabla = doc.add_table(rows=1, cols=2)
        tabla.autofit = False
        celda_texto, celda_foto = tabla.rows[0].cells
        celda_texto.width = Mm(130)
        celda_foto.width = Mm(30)
        destino = celda_texto
        # Toda celda nace con un párrafo vacío: se reutiliza en vez de añadir
        # otro, o el nombre saldría con una línea en blanco encima.
        primer_parrafo = celda_texto.paragraphs[0]
        try:
            celda_foto.paragraphs[0].add_run().add_picture(io.BytesIO(foto), width=Mm(28))
        except Exception:  # noqa: BLE001
            # Formato que python-docx no reconoce · el CV sale sin foto y ya.
            logger.warning("No se pudo incrustar la foto en el .docx")

    p = primer_parrafo if primer_parrafo is not None else destino.add_paragraph()
    run = p.add_run(cv.student_name)
    run.bold = True
    run.font.size = _pt(24)
    run.font.color.rgb = _rgb(_NAVY)
    p.paragraph_format.space_after = _pt(0)

    if cv.headline:
        ph = destino.add_paragraph()
        rh = ph.add_run(cv.headline)
        rh.font.size = _pt(11.5)
        rh.font.color.rgb = _rgb(_GRIS)
        rh.bold = True
        ph.paragraph_format.space_after = _pt(2)

    contacto = [
        x
        for x in (
            cv.current_occupation,
            cv.email,
            cv.school_name,
            f"Grado: {cv.grade}" if cv.grade else None,
            # Donde los idiomas van en sección propia (España, Europass) el
            # encabezado no los repite · mismo criterio que en el PDF.
            f"Inglés: {cv.english_level}"
            if (cv.english_level and not omitir_idioma)
            else None,
        )
        if x
    ]
    if contacto:
        pc = destino.add_paragraph()
        rc = pc.add_run("  ·  ".join(str(x) for x in contacto))
        rc.font.size = _pt(9.5)
        rc.font.color.rgb = _rgb(_GRIS)


def render_cv_docx(
    cv: CVData,
    *,
    estandar: str = ESTANDAR_POR_DEFECTO,
    estilo: str = ESTILO_POR_DEFECTO,
    incluir_foto: bool = True,
) -> bytes:
    """Devuelve los bytes del `.docx`.

    Misma firma que `render_cv_pdf` a propósito: el endpoint elige el
    renderizador por el parámetro `formato` y no tiene que saber que uno lleva
    CSS y el otro no.

    `estilo` se acepta y se ignora salvo por el tamaño base: en Word la
    apariencia la termina de decidir quien abre el archivo, y prometer tres
    diseños fieles al PDF sería prometer algo que el formato no sostiene.
    """
    try:
        from docx import Document
        from docx.shared import Mm
    except ImportError as exc:  # pragma: no cover
        raise RuntimeError(
            "python-docx no está instalado · agregá `python-docx==1.1.2` a requirements.txt"
        ) from exc

    est = obtener_estandar(estandar)
    sty = obtener_estilo(estilo)
    con_foto = debe_incluir_foto(
        est, quiere_foto=incluir_foto, hay_foto=bool(cv.photo_data_uri)
    )

    doc = Document()

    seccion = doc.sections[0]
    # El estándar de una página aprieta márgenes; los de dos respiran.
    margen = Mm(14) if est.max_paginas <= 1 else Mm(18)
    seccion.top_margin = seccion.bottom_margin = margen
    seccion.left_margin = seccion.right_margin = margen

    base = doc.styles["Normal"]
    base.font.name = "Calibri"
    base.font.size = _pt(9.5 if sty.clave == "compacto" else 10.5)

    _encabezado(doc, cv, con_foto=con_foto, omitir_idioma=idioma_va_en_seccion(est))

    for clave in est.orden_secciones:
        if clave in _SECCIONES:
            _SECCIONES[clave](doc, cv)

    # La cláusula legal del destino (RGPD en España) va antes del pie de
    # Mentoring: es parte del documento de la persona, no de nuestra firma.
    if est.aviso_legal:
        legal = doc.add_paragraph()
        rl = legal.add_run(est.aviso_legal)
        rl.font.size = _pt(8.5)
        rl.font.color.rgb = _rgb(_GRIS)
        legal.paragraph_format.space_before = _pt(10)

    pie = doc.add_paragraph()
    rp = pie.add_run(
        f"Hoja de Vida generada con Mentoring · {cv.generated_on} · "
        "documento personal del estudiante."
    )
    rp.font.size = _pt(8)
    rp.font.color.rgb = _rgb((0x9A, 0xA0, 0xAB))
    pie.paragraph_format.space_before = _pt(12)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
